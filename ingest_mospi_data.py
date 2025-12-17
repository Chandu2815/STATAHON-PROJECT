"""
Multi-format Data Ingestion Service for MoSPI Real Data
Supports: XLSX, DOCX, PDF files from microdata.gov.in
"""
import pandas as pd
import os
import sys
from pathlib import Path
import json
from typing import Dict, Any, List, Optional
import logging

# Document processing libraries
from openpyxl import load_workbook
from docx import Document
import pdfplumber
import PyPDF2

from app.database import SessionLocal
from app.models.dataset import Dataset, DataRecord
from app.config import get_settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

settings = get_settings()


class MoSPIDataIngestion:
    """Ingestion service for multiple file formats from MoSPI"""
    
    def __init__(self, data_dir: str = "data/mospi_real_data"):
        self.data_dir = Path(data_dir)
        self.db = SessionLocal()
        self.supported_extensions = {'.xlsx', '.xls', '.docx', '.pdf', '.csv'}
    
    def scan_directory(self) -> Dict[str, List[Path]]:
        """Scan directory and categorize files by type"""
        files_by_type = {
            'xlsx': [],
            'docx': [],
            'pdf': [],
            'csv': []
        }
        
        if not self.data_dir.exists():
            logger.warning(f"Directory {self.data_dir} does not exist")
            return files_by_type
        
        for file_path in self.data_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                ext = file_path.suffix.lower().replace('.', '')
                if ext == 'xls':
                    ext = 'xlsx'  # Treat XLS as XLSX
                if ext in files_by_type:
                    files_by_type[ext].append(file_path)
        
        return files_by_type
    
    def extract_xlsx_data(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from XLSX file"""
        logger.info(f"Processing XLSX: {file_path.name}")
        
        try:
            # Read with pandas for easy data manipulation
            excel_file = pd.ExcelFile(file_path)
            
            result = {
                'filename': file_path.name,
                'sheets': {},
                'metadata': {
                    'sheet_names': excel_file.sheet_names,
                    'total_sheets': len(excel_file.sheet_names)
                }
            }
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Clean data
                    df = df.dropna(how='all')  # Remove empty rows
                    df = df.dropna(axis=1, how='all')  # Remove empty columns
                    
                    if not df.empty:
                        result['sheets'][sheet_name] = {
                            'data': df.to_dict('records'),
                            'columns': list(df.columns),
                            'row_count': len(df),
                            'column_count': len(df.columns)
                        }
                        
                        logger.info(f"  Sheet '{sheet_name}': {len(df)} rows, {len(df.columns)} columns")
                except Exception as e:
                    logger.warning(f"  Could not read sheet '{sheet_name}': {e}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            return None
    
    def extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata and text from DOCX file"""
        logger.info(f"Processing DOCX: {file_path.name}")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables_data.append(table_data)
            
            result = {
                'filename': file_path.name,
                'paragraphs': paragraphs[:100],  # First 100 paragraphs
                'paragraph_count': len(paragraphs),
                'tables': tables_data,
                'table_count': len(tables_data),
                'metadata': {
                    'has_tables': len(tables_data) > 0,
                    'total_paragraphs': len(paragraphs)
                }
            }
            
            logger.info(f"  Extracted {len(paragraphs)} paragraphs, {len(tables_data)} tables")
            return result
            
        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            return None
    
    def extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata and text from PDF file"""
        logger.info(f"Processing PDF: {file_path.name}")
        
        try:
            result = {
                'filename': file_path.name,
                'pages': [],
                'metadata': {}
            }
            
            # Try pdfplumber first (better for tables)
            with pdfplumber.open(file_path) as pdf:
                result['metadata']['page_count'] = len(pdf.pages)
                
                # Extract text from first few pages
                max_pages = min(10, len(pdf.pages))
                for i, page in enumerate(pdf.pages[:max_pages]):
                    page_data = {
                        'page_number': i + 1,
                        'text': page.extract_text(),
                        'tables': page.extract_tables()
                    }
                    result['pages'].append(page_data)
            
            logger.info(f"  Extracted {len(result['pages'])} pages from PDF")
            return result
            
        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            return None
    
    def extract_csv_data(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from CSV file"""
        logger.info(f"Processing CSV: {file_path.name}")
        
        try:
            df = pd.read_csv(file_path)
            
            result = {
                'filename': file_path.name,
                'data': df.to_dict('records'),
                'columns': list(df.columns),
                'metadata': {
                    'row_count': len(df),
                    'column_count': len(df.columns)
                }
            }
            
            logger.info(f"  Extracted {len(df)} rows, {len(df.columns)} columns")
            return result
            
        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            return None
    
    def ingest_to_database(self, extracted_data: Dict[str, Any], dataset_name: str):
        """Save extracted data to database"""
        try:
            # Create or get dataset
            dataset = self.db.query(Dataset).filter(Dataset.name == dataset_name).first()
            
            if not dataset:
                dataset = Dataset(
                    name=dataset_name,
                    description=f"Data from {extracted_data.get('filename', 'unknown')}",
                    table_name=f"data_{dataset_name.lower().replace(' ', '_')}",
                    config=extracted_data.get('metadata', {})
                )
                self.db.add(dataset)
                self.db.commit()
                self.db.refresh(dataset)
                logger.info(f"Created dataset: {dataset_name}")
            
            # Insert data records
            if 'sheets' in extracted_data:
                # XLSX data with sheets
                for sheet_name, sheet_data in extracted_data['sheets'].items():
                    if 'data' in sheet_data and sheet_data['data']:
                        records = [
                            DataRecord(
                                dataset_id=dataset.id,
                                data={
                                    'sheet': sheet_name,
                                    'source_file': extracted_data['filename'],
                                    **record
                                }
                            )
                            for record in sheet_data['data'][:1000]  # Limit to 1000 records per sheet
                        ]
                        self.db.bulk_save_objects(records)
                        logger.info(f"  Inserted {len(records)} records from sheet '{sheet_name}'")
            
            elif 'data' in extracted_data:
                # CSV data
                records = [
                    DataRecord(
                        dataset_id=dataset.id,
                        data={
                            'source_file': extracted_data['filename'],
                            **record
                        }
                    )
                    for record in extracted_data['data'][:1000]  # Limit to 1000 records
                ]
                self.db.bulk_save_objects(records)
                logger.info(f"  Inserted {len(records)} records")
            
            self.db.commit()
            logger.info(f"✓ Successfully ingested data for dataset: {dataset_name}")
            
        except Exception as e:
            logger.error(f"Error ingesting to database: {e}")
            self.db.rollback()
    
    def process_all_files(self, save_to_db: bool = True):
        """Process all files in the data directory"""
        logger.info(f"Scanning directory: {self.data_dir}")
        
        files_by_type = self.scan_directory()
        
        summary = {
            'xlsx_files': len(files_by_type['xlsx']),
            'docx_files': len(files_by_type['docx']),
            'pdf_files': len(files_by_type['pdf']),
            'csv_files': len(files_by_type['csv']),
            'processed': [],
            'failed': []
        }
        
        print("\n" + "="*80)
        print(f"  FOUND FILES IN {self.data_dir}")
        print("="*80)
        print(f"  XLSX files: {summary['xlsx_files']}")
        print(f"  DOCX files: {summary['docx_files']}")
        print(f"  PDF files: {summary['pdf_files']}")
        print(f"  CSV files: {summary['csv_files']}")
        print("="*80 + "\n")
        
        # Process XLSX files
        for xlsx_file in files_by_type['xlsx']:
            try:
                data = self.extract_xlsx_data(xlsx_file)
                if data and save_to_db:
                    dataset_name = f"PLFS_{xlsx_file.stem}"
                    self.ingest_to_database(data, dataset_name)
                summary['processed'].append(xlsx_file.name)
            except Exception as e:
                logger.error(f"Failed to process {xlsx_file.name}: {e}")
                summary['failed'].append(xlsx_file.name)
        
        # Process CSV files
        for csv_file in files_by_type['csv']:
            try:
                data = self.extract_csv_data(csv_file)
                if data and save_to_db:
                    dataset_name = f"Data_{csv_file.stem}"
                    self.ingest_to_database(data, dataset_name)
                summary['processed'].append(csv_file.name)
            except Exception as e:
                logger.error(f"Failed to process {csv_file.name}: {e}")
                summary['failed'].append(csv_file.name)
        
        # Process DOCX files (metadata only)
        for docx_file in files_by_type['docx']:
            try:
                data = self.extract_docx_metadata(docx_file)
                summary['processed'].append(docx_file.name)
                # Save metadata to JSON file
                output_file = self.data_dir / f"{docx_file.stem}_metadata.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2)
                logger.info(f"  Saved metadata to {output_file.name}")
            except Exception as e:
                logger.error(f"Failed to process {docx_file.name}: {e}")
                summary['failed'].append(docx_file.name)
        
        # Process PDF files (metadata only)
        for pdf_file in files_by_type['pdf']:
            try:
                data = self.extract_pdf_metadata(pdf_file)
                summary['processed'].append(pdf_file.name)
                # Save metadata to JSON file
                output_file = self.data_dir / f"{pdf_file.stem}_metadata.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, default=str)
                logger.info(f"  Saved metadata to {output_file.name}")
            except Exception as e:
                logger.error(f"Failed to process {pdf_file.name}: {e}")
                summary['failed'].append(pdf_file.name)
        
        return summary
    
    def __del__(self):
        if self.db:
            self.db.close()


def main():
    """CLI entry point"""
    import argparse
    
    parser = argparse.ArgumentParser(description='MoSPI Multi-format Data Ingestion')
    parser.add_argument('--data-dir', default='data/mospi_real_data', help='Data directory path')
    parser.add_argument('--no-db', action='store_true', help='Skip database ingestion')
    parser.add_argument('--copy-from-downloads', action='store_true', help='Copy PLFS files from Downloads')
    
    args = parser.parse_args()
    
    # Copy files from Downloads if requested
    if args.copy_from_downloads:
        import shutil
        downloads = Path.home() / 'Downloads'
        target_dir = Path(args.data_dir)
        target_dir.mkdir(parents=True, exist_ok=True)
        
        plfs_files = [
            'Data_LayoutPLFS_Calendar_2024 (4).xlsx',
            'PLFS Panel 4 Sch 10.4 Item Code Description & Codes (1).xlsx',
            'District_codes_PLFS_Panel_4_202324_2024 (1).xlsx',
            'NMDS_2.0_PLFS_final upd (1).docx',
            'README_Calendar_2024 (3).docx',
            'Instruction Manual PLFS Vol-II (2).pdf',
            'Instruction manual PLFS Vol I (1).pdf'
        ]
        
        print("\nCopying PLFS files from Downloads...")
        for filename in plfs_files:
            src = downloads / filename
            if src.exists():
                dst = target_dir / filename
                shutil.copy2(src, dst)
                print(f"✓ Copied: {filename}")
            else:
                print(f"✗ Not found: {filename}")
    
    # Run ingestion
    print("\n" + "="*80)
    print("  MoSPI DATA INGESTION SERVICE")
    print("="*80 + "\n")
    
    ingestion = MoSPIDataIngestion(data_dir=args.data_dir)
    summary = ingestion.process_all_files(save_to_db=not args.no_db)
    
    print("\n" + "="*80)
    print("  INGESTION SUMMARY")
    print("="*80)
    print(f"  Successfully processed: {len(summary['processed'])} files")
    print(f"  Failed: {len(summary['failed'])} files")
    if summary['failed']:
        print(f"\n  Failed files:")
        for f in summary['failed']:
            print(f"    - {f}")
    print("="*80 + "\n")
    
    print("✓ Data ingestion complete!")
    print(f"📁 Check {args.data_dir} for metadata JSON files")
    if not args.no_db:
        print("💾 Check database for ingested records")


if __name__ == "__main__":
    main()
